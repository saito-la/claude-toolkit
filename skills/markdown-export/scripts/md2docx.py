#!/usr/bin/env python3
"""
md2docx.py — markdown → docx 変換 + 表の列幅自動調整

Usage:
    python3 md2docx.py input.md [-o output.docx] [--template gothic|default]

列幅の決定方法:
    各列のセルテキストを走査し、CJK文字を2文字換算した「表示幅」の最大値に
    比例した幅を割り当てる。最小幅（10%）を保証して潰れを防ぐ。
"""

import argparse
import subprocess
import sys
import unicodedata
from pathlib import Path

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
TEMPLATES = {
    "gothic": TEMPLATE_DIR / "reference-gothic.docx",
    "default": TEMPLATE_DIR / "reference-default.docx",
}
MIN_COL_RATIO = 0.08  # 1列あたり最低8%


def display_width(text: str) -> int:
    """CJK文字を2、それ以外を1として表示幅を推定する"""
    w = 0
    for ch in text:
        if unicodedata.east_asian_width(ch) in ("W", "F"):
            w += 2
        else:
            w += 1
    return w


def auto_adjust_tables(docx_path: Path, page_numbers: bool = False, fit_title: bool = False) -> None:
    """表の枠線付与＋列幅自動調整。任意でページ番号フッター・タイトル1行化も行う。"""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document(str(docx_path))
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin

    def emu_to_twips(emu: int) -> int:
        return int(emu * 1440 / 914400)

    def set_cell_width(cell, emu: int) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = etree.SubElement(tcPr, qn("w:tcW"))
        tcW.set(qn("w:w"), str(emu_to_twips(emu)))
        tcW.set(qn("w:type"), "dxa")

    def set_table_borders(table, color: str = "595959", sz: int = 6) -> None:
        """全テーブルに細い枠線（既定: 濃い灰色・0.75pt）を付与する。
        多くのテンプレートは表スタイルの枠線が透明なため、明示的に上書きする。"""
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.Element(qn("w:tblPr"))
            tbl.insert(0, tblPr)
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        borders = etree.Element(qn("w:tblBorders"))
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = etree.SubElement(borders, qn("w:" + edge))
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        # OOXML のスキーマ順序: tblBorders は shd/tblLayout/tblCellMar/tblLook より前
        anchor = None
        for tag in ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"):
            anchor = tblPr.find(qn(tag))
            if anchor is not None:
                break
        if anchor is not None:
            anchor.addprevious(borders)
        else:
            tblPr.append(borders)

    for table in doc.tables:
        set_table_borders(table)
        n_cols = len(table.columns)
        # 各列の最大表示幅を計算
        max_widths = [0] * n_cols
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < n_cols:
                    w = display_width(cell.text.strip())
                    if w > max_widths[j]:
                        max_widths[j] = w

        # 最小幅を保証して比率を正規化
        total_raw = sum(max_widths) or 1
        ratios = [max(w / total_raw, MIN_COL_RATIO) for w in max_widths]
        total_ratio = sum(ratios)
        ratios = [r / total_ratio for r in ratios]

        # 幅を適用
        col_emus = [int(usable * r) for r in ratios]
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < n_cols:
                    set_cell_width(cell, col_emus[j])

    # 任意: タイトル（先頭段落）を太字＋用紙幅に収まるサイズへ縮小（1行化）
    if fit_title and doc.paragraphs and doc.paragraphs[0].text.strip():
        import math
        from docx.shared import Pt
        title = doc.paragraphs[0]
        usable_pt = usable / 12700.0  # EMU→pt
        full = sum(2 if unicodedata.east_asian_width(c) in ("W", "F", "A") else 1
                   for c in title.text) / 2.0
        size = min(15, max(8, math.floor(usable_pt * 0.95 / full))) if full else 15
        for run in title.runs:
            run.font.bold = True
            run.font.size = Pt(size)

    # 任意: 全ページのフッター中央にページ番号（PAGEフィールド）。
    # テンプレの「先頭ページ別」「偶数奇数別」を無効化し、既定フッターを全ページに適用する。
    if page_numbers:
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        st = doc.settings.element
        eao = st.find(qn("w:evenAndOddHeaders"))
        if eao is not None:
            st.remove(eao)
        for s in doc.sections:
            s.different_first_page_header_footer = False
            fp = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
            for r in list(fp.runs):
                r._r.getparent().remove(r._r)
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
            run._r.append(b); run._r.append(it); run._r.append(e)

    doc.save(str(docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="markdown → docx 変換 + 表列幅自動調整")
    parser.add_argument("input", help="入力 markdown ファイル")
    parser.add_argument("-o", "--output", help="出力 docx パス（省略時は入力と同名）")
    parser.add_argument(
        "--template",
        choices=list(TEMPLATES.keys()),
        default="gothic",
        help="使用するテンプレート (default: gothic)",
    )
    parser.add_argument("--page-numbers", action="store_true",
                        help="全ページのフッター中央にページ番号を付ける")
    parser.add_argument("--fit-title", action="store_true",
                        help="タイトル（先頭段落）を太字にし用紙幅に収まるサイズへ縮小する")
    parser.add_argument("--lang", default="ja-JP",
                        help="文書言語（docxに記録。Google Docs変換後の既定言語になる） (default: ja-JP)")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"ERROR: {input_path} が見つかりません", file=sys.stderr)
        sys.exit(1)

    output_path = (
        Path(args.output).expanduser().resolve()
        if args.output
        else input_path.with_suffix(".docx")
    )
    template_path = TEMPLATES[args.template]

    # pandoc 変換
    cmd = [
        "pandoc",
        str(input_path),
        "-f",
        "markdown-auto_identifiers",  # 見出しIDを生成しない（docx→Google Docs でブックマークが付くのを防ぐ）
        "-M",
        f"lang={args.lang}",  # 文書言語をdocxに記録（Google Docs変換後の既定言語になる。既定 ja-JP）
        f"--reference-doc={template_path}",
        "-o",
        str(output_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"pandoc エラー:\n{result.stderr}", file=sys.stderr)
        sys.exit(result.returncode)

    # 後処理（表の枠線・列幅、任意でページ番号・タイトル1行化）
    auto_adjust_tables(output_path, page_numbers=args.page_numbers, fit_title=args.fit_title)

    # 文書言語を docx の docDefaults に記録（Google Docs 変換後の既定言語を決める。主言語 w:val も設定）
    from docx import Document as _Doc
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _El
    _d = _Doc(str(output_path))
    _dd = _d.styles.element.find(_qn("w:docDefaults"))
    _rprd = _dd.find(_qn("w:rPrDefault")) if _dd is not None else None
    _rpr = _rprd.find(_qn("w:rPr")) if _rprd is not None else None
    if _rpr is not None:
        _lang = _rpr.find(_qn("w:lang"))
        if _lang is None:
            _lang = _El("w:lang")
            _rpr.append(_lang)
        _lang.set(_qn("w:val"), args.lang)
        _lang.set(_qn("w:eastAsia"), args.lang)
        _d.save(str(output_path))

    print(f"✓ {output_path}")


if __name__ == "__main__":
    main()
